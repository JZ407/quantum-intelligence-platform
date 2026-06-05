"""
Streamlit App: 量子科技情报

Usage (local network):
    python -m streamlit run examples/daily_report_app.py --server.address 0.0.0.0 --server.port 8501
"""

import os
import sys
import io
import re
import json
from datetime import datetime, timedelta

import pandas as pd
import yaml
from sqlalchemy import create_engine
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# ------------------------------------------------------------------
# Config
# ------------------------------------------------------------------
DB_URL = 'mysql+pymysql://scraper:scraper123@127.0.0.1:3306/liangke_scraper?charset=utf8mb4'
HISTORICAL_DB_PATH = 'D:/Claude_code/liangke_historical/historical_final.db'
INSTITUTION_DB_PATH = 'D:/Claude_code/institution_news/institutions.db'
@st.cache_data(ttl=300)
def _get_inst_list():
    """Get distinct institution names from DB (cached 5 min)."""
    try:
        import sqlite3
        if not os.path.exists(INSTITUTION_DB_PATH):
            return []
        conn = sqlite3.connect(INSTITUTION_DB_PATH)
        c = conn.cursor()
        c.execute('SELECT DISTINCT source FROM articles ORDER BY source')
        result = [r[0] for r in c.fetchall()]
        conn.close()
        return result
    except Exception:
        return []
CATEGORY_PRIORITY = ['资本运作', '产品动态', '企业资讯', '科技前沿', '宏观态势']

CONF_DB_PATH = 'D:/Claude_code/conference_db/conferences.db'

TAGS_LIST = [
    '量子计算', '科技前沿', '产品动态', '量子通信', '行业应用',
    '企业与机构', '硬件平台', '融资商业', '宏观态势', 'AI/ML',
    '半导体', '量子物理', '后量子密码', '融资', 'PQC', 'QKD',
    '量子纠错', '超导', 'NIST', '量子传感', '企业资讯',
    '光量子', '资本运作', '离子阱', '政策标准', '后量子迁移', 'arXiv',
]

# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _set_run_font(run, font_name, size_pt):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)

def _parse_tags(x):
    if isinstance(x, list):
        return x
    if isinstance(x, str):
        try:
            return json.loads(x)
        except Exception:
            return []
    return []


def fetch_articles(target_date: str):
    """Fetch articles for a specific liangke_date."""
    engine = create_engine(DB_URL)
    query = f"SELECT * FROM articles WHERE liangke_date = '{target_date}' ORDER BY id DESC"
    df = pd.read_sql(query, engine)
    df['tags'] = df['tags'].apply(_parse_tags)
    return df


def _classify_by_title(title: str) -> list:
    """Rule-based classification using keywords in title.
    Returns exactly ONE tag from the 5 main categories.
    Priority: 资本运作 > 产品动态 > 企业资讯 > 科技前沿 > 宏观态势
    """
    t = title.lower()

    # 资本运作（最高优先级）
    cn_finance = ['融资', '投资', 'ipo', '并购', '资本', '轮', '美元', '亿元', '估值', '收购', '领投', '独角兽', '上市']
    en_finance = ['funding', 'fund', 'invest', 'merger', 'acquisition', 'capital', 'financing', 'valuation', 'raised', 'series a', 'series b', 'series c', 'seed round', '$m', '$b', 'million', 'billion', 'stake', 'acquire', 'acquires', 'investor', 'backing', 'venture', 'spin-off', 'spinoff', 'goes public', 'listed']
    if any(k in t for k in cn_finance + en_finance):
        return ['资本运作']

    # 产品动态
    cn_product = ['产品', '发布', '推出', '芯片', '计算机', '软件', '系统', '设备', '仪器', '平台', '上线', '原型机', '量子计算机', '量子芯片']
    en_product = ['product', 'launch', 'launches', 'release', 'releases', 'unveil', 'unveils', 'introduces', 'chip', 'computer', 'software', 'system', 'device', 'platform', 'processor', 'roadmap', 'road map', 'available now', 'now available', 'announces new', 'announced new', 'next-gen', 'next-generation', 'upgrade', 'debut', 'debuts', 'preview', 'beta', 'demo', 'demonstrated', 'demonstrates', 'showcases']
    if any(k in t for k in cn_product + en_product):
        return ['产品动态']

    # 企业资讯
    cn_biz = ['公司', '企业', '合作', '签约', '战略', '成立', '总部', '裁员', '人事', '任命', 'ceo', '总裁', '总监']
    en_biz = ['partner', 'partnership', 'collaboration', 'collaborate', 'agreement', 'appoints', 'appointed', 'named', 'joins', 'hires', 'leadership', 'executive', 'expands', 'expand', 'expansion', 'opens office', 'headquarters', 'company', 'enterprise', 'strategic', 'founded', 'president', 'director', 'alliance', 'consortium', 'mou', 'memorandum', 'team up', 'teams up', 'joint venture']
    if any(k in t for k in cn_biz + en_biz):
        return ['企业资讯']

    # 科技前沿
    cn_science = ['论文', '研究', '突破', '实验', '量子比特', '纠错', '算法', '物理', '科学', '发表', '期刊', '学术', '实验室', '原理', '理论']
    en_science = ['paper', 'research', 'researchers', 'breakthrough', 'experiment', 'qubit', 'qubits', 'algorithm', 'algorithms', 'physics', 'science', 'nature', 'published', 'journal', 'peer-reviewed', 'preprint', 'arxiv', 'discovery', 'discovered', 'achieved', 'record', 'milestone', 'novel', 'method', 'technique', 'fidelity', 'error correction', 'error mitigation', 'entanglement', 'coherence', 'superconducting', 'ion trap', 'neutral atom', 'photonic', 'topological', 'logical qubit', 'fault-tolerant', 'fault tolerant', 'simulation', 'simulator', 'advantage', 'supremacy']
    if any(k in t for k in cn_science + en_science):
        return ['科技前沿']

    # 宏观态势（兜底）
    return ['宏观态势']


def fetch_historical_articles(target_date: str):
    """Fetch articles from historical SQLite DB for a specific date."""
    import sqlite3
    conn = sqlite3.connect(HISTORICAL_DB_PATH)
    query = f"SELECT id, title, content, reference_url, liangke_url, liangke_date, tags FROM articles WHERE liangke_date LIKE '{target_date}%' ORDER BY liangke_date DESC"
    df = pd.read_sql(query, conn)
    conn.close()
    df['reference_url'] = df['liangke_url'].fillna(df['reference_url'])
    df['tags'] = df['tags'].apply(_parse_tags)
    return df


FINE_TAG_MAP = {
    '量子计算': ['quantum comput', 'quantum processor', 'quantum chip', 'quantum algorithm', 'quantum circuit',
                  'qubit', 'qubits', 'quantum advantage', 'quantum supremacy',
                  '量子计算', '量子计算机', '量子处理器', '量子芯片', '量子比特', '量子算法', '量子霸权', '量子优越'],
    '量子纠错': ['error correct', 'error mitigat', 'fault-tolerant', 'fault tolerant', 'logical qubit',
                  'surface code', 'decoherence', 'quantum error', 'fidelity',
                  '量子纠错', '纠错码', '容错', '逻辑量子比特', '退相干', '保真度'],
    '超导': ['superconducting', 'superconductor', 'transmon', 'josephson',
              '超导', '约瑟夫森'],
    '离子阱': ['ion trap', 'trapped ion', 'ytterbium', 'barium',
               '离子阱', '囚禁离子'],
    '光量子': ['photonic', 'photon', 'optical', 'squeezed light', '光量子', '光子'],
    '中性原子': ['neutral atom', 'rydberg', 'optical tweezer', '中性原子', '里德堡', '光镊'],
    '拓扑': ['topological', 'majorana', 'anyon', '拓扑', '马约拉纳'],
    'AI/ML': ['machine learning', 'deep learning', 'neural network', 'llm', 'gpt',
              'artificial intelligence', 'ai ', ' ai', 'ai-powered', 'transformer',
              '人工智能', '机器学习', '深度学习', '神经网络', '大模型'],
    '量子通信': ['quantum communic', 'qkd', 'quantum key', 'quantum network', 'quantum internet',
                  'teleportation', '量子通信', '量子密钥', '量子网络', '量子互联网', '隐形传态'],
    '量子传感': ['quantum sens', 'magnetometer', 'gravimeter', '量子传感', '磁力计', '重力仪'],
    '融资商业': ['funding', 'fund', 'invest', 'series a', 'series b', 'series c', 'raised',
                  'million', 'billion', 'ipo', 'venture', 'capital', 'valuation', 'acquire',
                  'acquisition', 'merger', 'stake', 'backing', 'spin-off', 'listed',
                  '融资', '投资', '亿元', '并购', '收购', '上市', '独角兽', '估值'],
    '后量子密码': ['post-quantum', 'pqc', 'cryptograph', 'encryption', 'nist',
                   '后量子密码', '量子密码', '加密', '抗量子'],
    '政策标准': ['policy', 'regulation', 'standard', 'framework', 'initiative', 'government',
                  '政策', '标准', '框架', '政府', '倡议', '监管'],
    '半导体': ['semiconductor', 'cmos', 'fabrication', 'foundry', 'wafer',
               '半导体', '芯片制造', '晶圆', '代工'],
    '产品动态': ['product', 'launch', 'launches', 'release', 'unveil', 'unveils', 'introduces',
                  'processor', 'roadmap', 'road map', 'sdk', 'cloud', 'available now',
                  'now available', 'next-gen', 'next generation', 'debut', 'upgrade', 'beta',
                  '产品', '发布', '推出', '上线', '平台', '软件', '系统', '设备', '原型机'],
    '企业资讯': ['partner', 'partnership', 'collaborat', 'alliance', 'consortium', 'joint venture',
                  'appoints', 'appointed', 'named', 'joins', 'hires', 'ceo', 'president',
                  'executive', 'headquarters', 'expands', 'expand', 'expansion',
                  '企业', '公司', '合作', '战略', '任命', '总部', '成立'],
    '科技前沿': ['research', 'researchers', 'breakthrough', 'paper', 'published', 'journal',
                  'nature', 'science', 'discovery', 'milestone', 'record', 'novel',
                  'method', 'technique', 'experiment', 'physics', 'simulation',
                  '论文', '研究', '突破', '实验', '发表', '期刊', '学术', '物理', '科学'],
    '宏观态势': ['market', 'industry', 'report', 'forecast', 'trend', 'outlook',
                  '市场', '行业', '报告', '预测', '趋势', '前景'],
}


def _classify_inst_tags(title: str) -> list:
    """Fine-grained multi-tag classification for institution news (EN/CN)."""
    t = title.lower()
    tags = []
    for tag, keywords in FINE_TAG_MAP.items():
        if any(k in t for k in keywords):
            tags.append(tag)
    if not tags:
        tags.append('宏观态势')
    return tags


def fetch_institution_articles():
    """Fetch all articles from institution news DB."""
    import sqlite3
    if not os.path.exists(INSTITUTION_DB_PATH):
        return pd.DataFrame()
    conn = sqlite3.connect(INSTITUTION_DB_PATH)
    df = pd.read_sql("SELECT id, title, content, url, source, publish_date, tags, summary, title_cn FROM articles ORDER BY CASE WHEN publish_date IS NULL OR publish_date = '' THEN 1 ELSE 0 END, publish_date DESC, id DESC", conn)
    conn.close()
    df = df.rename(columns={'publish_date': 'liangke_date'})
    df['tags'] = df['title'].apply(_classify_inst_tags)
    df['source_tag'] = df['source']
    return df


def select_top3(df: pd.DataFrame):
    """Select top-3 articles by category priority."""
    if df.empty:
        return []

    selected = []
    used_ids = set()

    for cat in CATEGORY_PRIORITY:
        mask = df['tags'].apply(lambda tags: isinstance(tags, list) and cat in tags)
        candidates = df[mask]
        if not candidates.empty:
            # Sort by fetch_count desc, then id desc (newer first)
            candidates = candidates.sort_values(by=['fetch_count', 'id'], ascending=[False, False])
            for _, row in candidates.iterrows():
                if row['id'] not in used_ids:
                    selected.append(row)
                    used_ids.add(row['id'])
                    break
        if len(selected) >= 3:
            break

    # If still < 3, fill with remaining articles (newest first)
    if len(selected) < 3:
        remaining = df[~df['id'].isin(used_ids)].sort_values(by='id', ascending=False)
        for _, row in remaining.iterrows():
            selected.append(row)
            if len(selected) >= 3:
                break

    return selected

def build_docx(date_str: str, articles: list) -> io.BytesIO:
    """Generate a Word document and return as BytesIO."""
    doc = Document()
    YAHEI = '微软雅黑'

    # Title
    p = doc.add_paragraph()
    run = p.add_run(f"每日情报资讯（{date_str}）：")
    _set_run_font(run, YAHEI, 16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph()  # blank line

    for idx, art in enumerate(articles, 1):
        # Article title
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}、{art['title']}：")
        _set_run_font(run, YAHEI, 16)
        run.font.bold = True

        # Full content (split by newline, indent each paragraph)
        content = art.get('content', '') or ''
        for line in content.strip().split('\n'):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)  # ~2 Chinese chars
            run = p.add_run(line)
            _set_run_font(run, YAHEI, 14)

        # Reference link
        ref_url = art.get('reference_url', '') or art.get('liangke_url', '')
        if ref_url:
            p = doc.add_paragraph()
            run = p.add_run(f"参考链接：{ref_url}")
            _set_run_font(run, YAHEI, 10.5)
            run.font.color.rgb = RGBColor(0, 0, 255)

        doc.add_paragraph()  # blank line between articles

    # Patent section
    p = doc.add_paragraph()
    run = p.add_run("4、专利：")
    _set_run_font(run, YAHEI, 16)
    run.font.bold = True

    p = doc.add_paragraph()
    run = p.add_run(" ")
    _set_run_font(run, YAHEI, 14)

    p = doc.add_paragraph()
    run = p.add_run("参考链接：")
    _set_run_font(run, YAHEI, 10.5)
    run.font.color.rgb = RGBColor(0, 0, 255)

    # Save to memory
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ------------------------------------------------------------------
# Conference helpers
# ------------------------------------------------------------------

@st.cache_data(ttl=60)
def load_conferences():
    """Load translated conferences from SQLite database."""
    import sqlite3
    if not os.path.exists(CONF_DB_PATH):
        return []
    conn = sqlite3.connect(CONF_DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT date_str, month, name_zh, location_zh, url FROM conferences ORDER BY month, id')
    rows = [dict(r) for r in cursor.fetchall()]
    conn.close()
    return rows


# ------------------------------------------------------------------
# Article detail dialog
# ------------------------------------------------------------------

@st.dialog("新闻详情", width="large")
def show_article_detail(art: dict):
    """Show article detail in a modal dialog."""
    st.subheader(art['title'])
    tags = art.get('tags', [])
    inst = art.get('source', '') or art.get('source_tag', '')
    caption_parts = []
    if inst:
        caption_parts.append(f"机构：{inst}")
    if isinstance(tags, dict):
        weekly = tags.get('weekly', [])
        search = tags.get('search_tags', [])
        if weekly:
            caption_parts.append(f"周报：{' | '.join(weekly)}")
        if search:
            caption_parts.append(f"检索：{' | '.join(search[:5])}")
    elif isinstance(tags, list) and tags:
        caption_parts.append(f"标签：{' | '.join(tags)}")
    if caption_parts:
        st.caption(' · '.join(caption_parts))
    st.markdown("---")
    title_cn = art.get('title_cn', '') or ''
    if title_cn:
        st.markdown(f"**中文标题**：{title_cn}")
    content = art.get('content', '') or ''
    if content:
        with st.expander("原文内容", expanded=False):
            st.markdown(content)
    else:
        st.info("暂无正文内容")
    ref_url = art.get('url', '') or art.get('reference_url', '') or art.get('liangke_url', '')
    if ref_url:
        st.markdown("---")
        st.link_button("🔗 打开参考链接", ref_url)


# ------------------------------------------------------------------
# Page: Daily News
# ------------------------------------------------------------------

def page_daily_news():
    """Daily news selection and report generation."""
    st.header("量科网每日情报资讯")

    today = datetime.now().date()
    col1, col2, col3 = st.columns([2, 2, 2])
    with col1:
        source = st.selectbox("数据源", options=["量科每日库", "量科历史库", "机构新闻库"], index=0)
    with col2:
        if source == "量科每日库":
            min_date = datetime(2026, 4, 11).date()
            caption = "📌 每日库记录始于 2026-04-11"
        elif source == "量科历史库":
            min_date = datetime(2021, 11, 18).date()
            caption = "📌 历史库记录始于 2021-11-18"
        else:
            min_date = datetime(2020, 1, 1).date()
            caption = "📌 IBM/Quantinuum/NVIDIA/Google 官方新闻"
        target_date = st.date_input("选择日期", value=today, min_value=min_date)
    with col3:
        if source == "机构新闻库":
            inst_list = _get_inst_list()
            inst_filter = st.selectbox("机构筛选", options=["全部"] + inst_list, index=0)
        else:
            inst_filter = "全部"  # dummy, not used for other sources
        st.caption(caption)
    target_str = target_date.strftime('%Y-%m-%d')

    # Page type filter (only for 量科 sources)
    if source != "机构新闻库":
        page_types = st.multiselect(
            "内容类型", options=["flash", "article", "reference"],
            default=["flash", "article", "reference"], key="page_type_filter"
        )
    else:
        page_types = ["flash", "article", "reference"]

    keyword = st.text_input("🔍 关键词检索（搜索全库，不限日期）", placeholder="输入关键词搜索全库...")

    with st.spinner("正在读取数据库..."):
        if source == "机构新闻库":
            df = fetch_institution_articles()
            if inst_filter != "全部":
                df = df[df['source'] == inst_filter]
            if keyword:
                kw = keyword.strip()
                mask = df['title'].str.contains(kw, case=False, na=False) | df['content'].str.contains(kw, case=False, na=False)
                df = df[mask]
        elif keyword:
            # Search entire database, ignore date
            if source == "量科历史库":
                df = fetch_historical_articles_range('2021-01-01', '2030-01-01')
            else:
                df = _fetch_daily_articles_range('2026-01-01', '2030-01-01')
            kw = keyword.strip()
            mask = df['title'].str.contains(kw, case=False, na=False) | df['content'].str.contains(kw, case=False, na=False)
            df = df[mask]
        else:
            if source == "量科历史库":
                df = fetch_historical_articles(target_str)
            else:
                df = fetch_articles(target_str)

    # Apply page type filter (skip if all types selected = no filter)
    ALL_TYPES = {"flash", "article", "reference"}
    if set(page_types) != ALL_TYPES and not df.empty and 'page_type' in df.columns:
        df = df[df['page_type'].isin(page_types)]

    if df.empty:
        st.warning(f"📭 暂无匹配数据。")
        return

    # News list section (with manual selection)
    filter_note = "" if set(page_types) == ALL_TYPES else f" [类型: {', '.join(page_types)}]"
    if keyword:
        st.markdown(f"### 📋 搜索结果（共 {len(df)} 条）{filter_note}")
        st.caption(f"全文检索 \"{keyword}\" ，跨全库")
    else:
        st.markdown(f"### 📋 新闻列表（{target_str}，共 {len(df)} 条）{filter_note}")
        if source == "量科每日库":
            st.caption("请勾选您认为最重要的 3 条新闻，下方将据此生成日报。")

    # Data export (collapsible, placed above news list)
    with st.expander("📤 数据导出", expanded=False):
        exp_col1, exp_col2 = st.columns(2)
        with exp_col1:
            exp_start = st.date_input("开始日期", value=target_date, min_value=min_date, key="exp_start")
        with exp_col2:
            exp_end = st.date_input("结束日期", value=target_date, min_value=min_date, key="exp_end")
        exp_kw = st.text_input("关键词筛选", placeholder="留空=全部", key="exp_kw")
        exp_tags = st.multiselect("标签筛选（留空=全部）", options=TAGS_LIST, default=[], key="exp_tags")
        exp_format = st.radio("导出格式", ["Excel (.xlsx)", "SQLite (.db)"], horizontal=True, key="exp_format")

        if st.button("📤 导出数据", type="primary", key="btn_export"):
            with st.spinner("正在读取数据库..."):
                if source == "机构新闻库":
                    exp_df = fetch_institution_articles()
                    if inst_filter != "全部":
                        exp_df = exp_df[exp_df['source'] == inst_filter]
                elif source == "量科历史库":
                    exp_df = fetch_historical_articles_range(exp_start.strftime('%Y-%m-%d'), exp_end.strftime('%Y-%m-%d'))
                else:
                    exp_df = _fetch_daily_articles_range(exp_start.strftime('%Y-%m-%d'), exp_end.strftime('%Y-%m-%d'))
            if exp_kw:
                kw = exp_kw.strip()
                mask = exp_df['title'].str.contains(kw, case=False, na=False) | exp_df['content'].str.contains(kw, case=False, na=False)
                exp_df = exp_df[mask]
            if exp_tags:
                mask = exp_df['tags'].apply(lambda t: False if not t else (
                    isinstance(t, dict) and any(tag in t.get('weekly', []) + t.get('search_tags', []) for tag in exp_tags)
                    or (isinstance(t, list) and any(tag in t for tag in exp_tags))
                ))
                exp_df = exp_df[mask]
            if exp_df.empty:
                st.warning("未找到匹配的数据。")
            else:
                if exp_format == "Excel (.xlsx)":
                    buf = io.BytesIO()
                    exp_df.to_excel(buf, index=False, engine='openpyxl')
                    buf.seek(0)
                    st.download_button(
                        label="📥 下载 Excel",
                        data=buf,
                        file_name=f"export_{exp_start.strftime('%Y%m%d')}_{exp_end.strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )
                    st.success(f"已导出 {len(exp_df)} 条数据")
                else:
                    import sqlite3
                    tmp_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'export_tmp.db')
                    os.makedirs(os.path.dirname(tmp_path), exist_ok=True)
                    tmp_conn = sqlite3.connect(tmp_path)
                    exp_df.to_sql('articles', tmp_conn, index=False, if_exists='replace')
                    tmp_conn.close()
                    with open(tmp_path, 'rb') as f:
                        st.download_button(
                            label="📥 下载 SQLite 数据库",
                            data=f,
                            file_name=f"export_{exp_start.strftime('%Y%m%d')}_{exp_end.strftime('%Y%m%d')}.db",
                            mime="application/octet-stream",
                        )
                    st.success(f"已导出 {len(exp_df)} 条数据")

    if source == "量科历史库":
        source_key = "hist"
    elif source == "机构新闻库":
        source_key = "inst"
    else:
        source_key = "daily"

    # Page reset on source/keyword change
    page_key = f"page_{source_key}"
    trigger_key = f"trigger_{source_key}_{keyword}_{target_str}"
    if "last_trigger" not in st.session_state or st.session_state.get("last_trigger") != trigger_key:
        st.session_state[page_key] = 1
        st.session_state["last_trigger"] = trigger_key

    # Pagination
    PAGE_SIZE = 20
    total_pages = max(1, (len(df) + PAGE_SIZE - 1) // PAGE_SIZE)
    current_page = st.session_state.get(page_key, 1)

    if total_pages > 1:
        pc1, pc2, pc3 = st.columns([1, 2, 1])
        with pc1:
            if st.button("◀ 上一页", disabled=(current_page <= 1), key=f"prev_{source_key}"):
                st.session_state[page_key] = max(1, current_page - 1)
                st.rerun()
        with pc2:
            st.markdown(f"<div style='text-align:center;padding-top:5px'>第 {current_page} / {total_pages} 页</div>", unsafe_allow_html=True)
        with pc3:
            if st.button("下一页 ▶", disabled=(current_page >= total_pages), key=f"next_{source_key}"):
                st.session_state[page_key] = min(total_pages, current_page + 1)
                st.rerun()

    start_idx = (current_page - 1) * PAGE_SIZE
    df_page = df.iloc[start_idx:start_idx + PAGE_SIZE]

    show_checkbox = (source == "量科每日库")
    selected_ids = []
    for _, row in df_page.iterrows():
        with st.container():
            cols = st.columns(([0.5, 5, 1] if show_checkbox else [5, 1]))
            col_idx = 0
            if show_checkbox:
                with cols[0]:
                    checked = st.checkbox("", key=f"sel_{source_key}_{row['id']}", label_visibility="collapsed")
                    if checked:
                        selected_ids.append(row['id'])
                col_idx = 1
            with cols[col_idx]:
                art_url = row.get('url', '') or row.get('reference_url', '') or row.get('liangke_url', '')
                title_cn = row.get('title_cn', '') or ''
                if art_url:
                    st.markdown(f"[**{row['title']}**]({art_url})")
                else:
                    st.markdown(f"**{row['title']}**")
                if title_cn:
                    st.caption(title_cn)
                tags = row.get('tags', [])
                date_str = row.get('liangke_date', '')
                if date_str and len(str(date_str)) > 10:
                    date_str = str(date_str)[:10]
                # Handle both new dict format and old list format
                if isinstance(tags, dict):
                    tag_text = ' | '.join(tags.get('weekly', []))
                elif isinstance(tags, list) and tags:
                    tag_text = ' | '.join(tags[:3])
                else:
                    tag_text = ''
                inst = row.get('source', '') or row.get('source_tag', '')
                page_type = row.get('page_type', '')
                type_badge = f'[{page_type}] ' if page_type else ''
                if tag_text and date_str:
                    st.caption(f"{type_badge}{date_str} · {tag_text}{' · ' + inst if inst else ''}")
                elif date_str:
                    st.caption(f"{type_badge}{date_str}{' · ' + inst if inst else ''}")
                elif tag_text:
                    st.caption(f"{type_badge}{tag_text}{' · ' + inst if inst else ''}")
                elif inst:
                    st.caption(f"{type_badge}{inst}")
            detail_col = cols[2] if show_checkbox else cols[1]
            with detail_col:
                if st.button("查看详情", key=f"view_{source_key}_{row['id']}", type="secondary"):
                    show_article_detail(row.to_dict())
        st.divider()

    # Daily report generation section (only for daily news)
    if source == "量科每日库":
        st.markdown("---")
        st.markdown("### 📄 日报生成")

        selected_rows = df[df['id'].isin(selected_ids)].to_dict('records')
        count = len(selected_rows)

        if count == 0:
            st.info("请在上方新闻列表中勾选要纳入日报的文章（建议 3 条）。")
        elif count < 3:
            st.warning(f"已勾选 {count} 条，建议再选 {3 - count} 条以达到 3 条。")
        elif count > 3:
            st.warning(f"已勾选 {count} 条，建议只保留最重要的 3 条。当前将使用前 3 条生成日报。")
            selected_rows = selected_rows[:3]
        else:
            st.success(f"已勾选 {count} 条，可以生成日报。")

        if selected_rows and st.button("🚀 生成日报", type="primary"):
            rows_to_use = selected_rows[:3]

            st.markdown("#### 日报预览")
            for idx, art in enumerate(rows_to_use, 1):
                with st.container():
                    st.markdown(f"**{idx}. {art['title']}**")
                    tags = art.get('tags', [])
                    if isinstance(tags, dict):
                        weekly = tags.get('weekly', [])
                        search = tags.get('search_tags', [])
                        if weekly:
                            st.caption(f"周报：{' | '.join(weekly)}")
                        if search:
                            st.caption(f"检索：{' | '.join(search[:5])}")
                    elif isinstance(tags, list) and tags:
                        st.caption(f"标签：{' | '.join(tags)}")
                    st.markdown("---")

            doc_buf = build_docx(target_str, rows_to_use)
            file_name = f"日报{target_str}.docx"

            st.download_button(
                label="📥 下载 Word 日报",
                data=doc_buf,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )



def _fetch_daily_articles_range(start_date: str, end_date: str) -> pd.DataFrame:
    """Fetch articles from daily MySQL DB for a date range."""
    engine = create_engine(DB_URL)
    query = f"SELECT id, title, content, reference_url, liangke_url, liangke_date, tags FROM articles WHERE liangke_date BETWEEN '{start_date}' AND '{end_date}' ORDER BY id DESC"
    df = pd.read_sql(query, engine)
    df['tags'] = df['tags'].apply(_parse_tags)
    return df


def fetch_historical_articles_range(start_date: str, end_date: str) -> pd.DataFrame:
    """Fetch articles from historical SQLite DB for a date range."""
    import sqlite3
    conn = sqlite3.connect(HISTORICAL_DB_PATH)
    query = f"SELECT id, title, content, reference_url, liangke_date, tags FROM articles WHERE liangke_date >= '{start_date}' AND liangke_date < '{end_date} 23:59:59' ORDER BY liangke_date DESC"
    df = pd.read_sql(query, conn)
    conn.close()
    df['tags'] = df['tags'].apply(_parse_tags)
    return df


# ------------------------------------------------------------------
# Page: Conferences
# ------------------------------------------------------------------

def page_conferences():
    """Conference list viewer."""
    confs = load_conferences()
    if not confs:
        st.warning("📭 暂无会议数据。请先运行会议缓存构建脚本。")
        return

    st.markdown("### 📅 量子科技行业会议一览")
    st.caption(f"数据来源：quantum.info/conf/，共 {len(confs)} 场会议")

    month = st.selectbox(
        "选择月份",
        options=list(range(1, 13)),
        format_func=lambda x: f"{x}月",
        index=datetime.now().month - 1,
    )

    filtered = [c for c in confs if c.get('month') == month]

    if not filtered:
        st.info(f"{month} 月暂无会议信息。")
        return

    st.success(f"{month} 月共 {len(filtered)} 场会议")

    # Build display table
    rows = []
    for c in filtered:
        rows.append({
            '日期': c.get('date_str', ''),
            '会议名称': c.get('name_zh', c.get('name_en', '')),
            '地点': c.get('location_zh', c.get('location_en', '')),
            '链接': c.get('url', ''),
        })

    df = pd.DataFrame(rows)
    st.dataframe(df, use_container_width=True, hide_index=True)


# ------------------------------------------------------------------
# Page: Weekly Report
# ------------------------------------------------------------------

def _find_column(df: pd.DataFrame, candidates: list):
    cols_lower = {c.lower(): c for c in df.columns}
    for cand in candidates:
        if cand.lower() in cols_lower:
            return cols_lower[cand.lower()]
    return None


def parse_tender_excel(df: pd.DataFrame) -> list:
    col_name = _find_column(df, ['项目名称', '项目标题', 'title', 'name'])
    col_desc = _find_column(df, ['项目基本情况', '项目内容', 'desc', 'description'])
    col_unit = _find_column(df, ['项目主体单位', '采购单位', 'unit', 'organization'])
    col_scale = _find_column(df, ['项目规模', '预算金额', 'scale', 'budget'])
    col_pub = _find_column(df, ['发布时间', '发布日期', 'pub_date', 'publish date'])
    col_pre = _find_column(df, ['预采时间', '预计采购时间', 'pre_date'])
    col_url = _find_column(df, ['信息来源', '来源链接', 'url', 'source', 'link'])

    tenders = []
    for _, row in df.iterrows():
        t = {
            'name': str(row[col_name]) if col_name and pd.notna(row[col_name]) else '',
            'desc': str(row[col_desc]) if col_desc and pd.notna(row[col_desc]) else '',
            'unit': str(row[col_unit]) if col_unit and pd.notna(row[col_unit]) else '',
            'scale': str(row[col_scale]) if col_scale and pd.notna(row[col_scale]) else '',
            'pub_date': str(row[col_pub]) if col_pub and pd.notna(row[col_pub]) else '',
            'pre_date': str(row[col_pre]) if col_pre and pd.notna(row[col_pre]) else '',
            'url': str(row[col_url]) if col_url and pd.notna(row[col_url]) else '',
        }
        t = {k: (v if v not in ('nan', 'None', 'null') else '') for k, v in t.items()}
        tenders.append(t)
    return tenders


def parse_patent_excel(df: pd.DataFrame, file_bytes=None) -> list:
    col_title = _find_column(df, ['标题(译)(简体中文)', '发明名称(中文)(机器翻译)', '标题', 'title'])
    col_applicant = _find_column(df, ['[标]当前申请(专利权)人', '申请人', 'applicant'])
    col_inventor = _find_column(df, ['发明人', 'inventor'])
    col_type = _find_column(df, ['法律状态/事件', '法律状态', '专利类型', 'type'])
    col_date = _find_column(df, ['公开(公告)日', '公开日', '授权日', 'date'])
    col_abstract = _find_column(df, ['摘要(译)(简体中文)', '摘要(中文)(机器翻译)', '摘要', 'abstract'])

    # Extract hyperlinks from 公开号 column
    hyperlinks = {}
    if file_bytes:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes))
            ws = wb.active
            num_col_idx = None
            for col in range(1, ws.max_column + 1):
                v = ws.cell(row=1, column=col).value
                if v and '公开' in str(v) and '号' in str(v):
                    num_col_idx = col
                    break
            if num_col_idx:
                for row_idx in range(2, ws.max_row + 1):
                    cell = ws.cell(row=row_idx, column=num_col_idx)
                    if cell.hyperlink and cell.hyperlink.target:
                        hyperlinks[row_idx - 2] = cell.hyperlink.target
            wb.close()
        except Exception:
            pass

    patents = []
    for i, (_, row) in enumerate(df.iterrows()):
        ptype = str(row[col_type]) if col_type and pd.notna(row[col_type]) else ''
        if '授权' in ptype:
            ptype = '授权专利'
        elif '公开' in ptype:
            ptype = '公开专利'
        else:
            ptype = '公开专利'
        url = hyperlinks.get(i, '') or ''
        p = {
            'title': str(row[col_title]) if col_title and pd.notna(row[col_title]) else '',
            'applicant': str(row[col_applicant]) if col_applicant and pd.notna(row[col_applicant]) else '',
            'inventor': str(row[col_inventor]) if col_inventor and pd.notna(row[col_inventor]) else '',
            'type': ptype,
            'date': str(row[col_date]) if col_date and pd.notna(row[col_date]) else '',
            'abstract': str(row[col_abstract]) if col_abstract and pd.notna(row[col_abstract]) else '',
            'url': url,
        }
        p = {k: (v if v not in ('nan', 'None', 'null') else '') for k, v in p.items()}
        patents.append(p)
    return patents


def filter_patents_by_date(patents: list, start_date, end_date) -> list:
    """Filter patents to only those with dates within the week range."""
    from datetime import datetime as dt2
    result = []
    for p in patents:
        d = p.get('date', '')
        if not d:
            result.append(p)  # keep if no date
            continue
        try:
            # Try common date formats
            for fmt in ['%Y-%m-%d', '%Y%m%d', '%Y/%m/%d', '%Y年%m月%d日']:
                try:
                    pd_date = dt2.strptime(str(d).strip()[:10], fmt).date()
                    if pd_date >= start_date and pd_date <= end_date:
                        result.append(p)
                    break
                except ValueError:
                    continue
        except Exception:
            result.append(p)  # keep if can't parse
    return result


def filter_patents_llm(patents: list, category: str, max_keep: int = 8) -> list:
    """LLM ranks patents, prioritizing 授权 > 公开 > 实用新型 > 外观."""
    if len(patents) <= max_keep:
        return patents

    client = _get_weekly_llm()
    lines = [
        f"以下是【{category}】板块的{len(patents)}条专利。请筛选出最重要的{max_keep}条放入周报。",
        "优先级：授权专利 > 公开专利 > 实用新型专利 > 外观专利。在同类中，优先创新型、突破性技术。",
        "只输出选中的序号（逗号分隔），不要解释。\n"
    ]
    for i, p in enumerate(patents, 1):
        lines.append(f"{i}. [{p.get('type','未知')}] {p.get('title','')} | {p.get('applicant','')} | {p.get('abstract','')[:80]}")

    messages = [
        {"role": "system", "content": "你是专利分析专家。"},
        {"role": "user", "content": "\n".join(lines)},
    ]
    try:
        resp = client.chat(messages)
        nums = re.findall(r'\d+', resp)
        selected = []
        for n in nums:
            idx = int(n) - 1
            if 0 <= idx < len(patents) and patents[idx] not in selected:
                selected.append(patents[idx])
        return selected[:max_keep] if selected else patents[:max_keep]
    except Exception:
        return patents[:max_keep]


def _get_weekly_llm():
    sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'rag_system'))
    from llm_client import LLMClient
    cfg_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'config.yaml')
    cfg = yaml.safe_load(open(cfg_path, encoding='utf-8'))['llm']
    return LLMClient(
        provider='openai', api_key=cfg['api_key'], api_base=cfg['api_base'],
        model=cfg['model'], max_tokens=2048, timeout=180,
    )


def page_weekly_report():
    st.header("周报生成")

    # Default: this week's Saturday → today (current week in progress)
    today = datetime.now().date()
    # Week runs Saturday→Friday. Find the Saturday that starts this week.
    # Mon(0)→1, Tue(1)→2, Wed(2)→3, Thu(3)→4, Fri(4)→5, Sat(5)→0, Sun(6)→1
    wd = today.weekday()
    if wd == 5:       # Saturday → today IS the start
        days_back = 0
    elif wd == 6:     # Sunday → yesterday was the start
        days_back = 1
    else:             # Mon-Fri
        days_back = wd + 1
    last_saturday = today - timedelta(days=days_back)

    # Auto-calculate issue number: issue 182 = week starting 2026-05-31 (Sat)
    REF_ISSUE = 182
    REF_SATURDAY = datetime(2026, 5, 31).date()
    days_diff = (last_saturday - REF_SATURDAY).days
    weeks_diff = (days_diff + 6) // 7 if days_diff >= 0 else -((-days_diff + 6) // 7)
    default_issue = str(REF_ISSUE + weeks_diff)

    col1, col2, col3 = st.columns(3)
    with col1:
        start_date = st.date_input("开始日期", value=last_saturday, key="wr_start")
    with col2:
        end_date = st.date_input("结束日期", value=today, key="wr_end")
    with col3:
        issue_no = st.text_input("期数", value=default_issue, key="wr_issue")

    conf_month = st.selectbox("会议月份", options=list(range(1, 13)),
                               format_func=lambda x: f"{x}月",
                               index=datetime.now().month - 1, key="wr_conf")

    st.markdown("---")
    wr_page_types = st.multiselect(
        "新闻类型筛选（默认全选，叉掉不想要的）",
        options=["flash", "article", "reference"],
        default=["flash", "article", "reference"], key="wr_page_types"
    )

    # --- Article preview & selection ---
    st.markdown("---")
    st.markdown("### 新闻勾选")
    preview_col1, preview_col2 = st.columns([2, 5])
    with preview_col1:
        load_clicked = st.button("📋 加载文章预览", type="secondary", key="wr_load")

    if 'wr_selections' not in st.session_state:
        st.session_state.wr_selections = {}
    if 'wr_loaded' not in st.session_state:
        st.session_state.wr_loaded = False

    if load_clicked:
        with st.spinner("正在加载文章..."):
            engine = create_engine(DB_URL)
            start_str = start_date.strftime('%Y-%m-%d')
            end_str = end_date.strftime('%Y-%m-%d')
            query = f"SELECT id, title, liangke_date, page_type FROM articles WHERE liangke_date BETWEEN '{start_str}' AND '{end_str}' ORDER BY id DESC"
            df_raw = pd.read_sql(query, engine)
            if wr_page_types and 'page_type' in df_raw.columns:
                df_raw = df_raw[df_raw['page_type'].isin(wr_page_types)]
            st.session_state.wr_articles = df_raw
            st.session_state.wr_loaded = True
            # Initialize all as selected
            st.session_state.wr_selections = {str(row['id']): True for _, row in df_raw.iterrows()}
            # Also load tags for classification
            tag_query = f"SELECT id, tags FROM articles WHERE liangke_date BETWEEN '{start_str}' AND '{end_str}'"
            df_tags = pd.read_sql(tag_query, engine)
            df_tags['tags'] = df_tags['tags'].apply(_parse_tags)
            st.session_state.wr_tags = dict(zip(df_tags['id'].astype(str), df_tags['tags']))
        st.rerun()

    if st.session_state.wr_loaded and 'wr_articles' in st.session_state:
        df = st.session_state.wr_articles

        # Classify into 5 categories using tags
        CATS = ['资本运作', '产品动态', '企业资讯', '科技前沿', '宏观态势']
        categories = {cat: [] for cat in CATS}
        uncategorized = []
        for _, row in df.iterrows():
            aid = str(row['id'])
            tags = st.session_state.wr_tags.get(aid, [])
            matched = False
            if isinstance(tags, dict):
                for tag in tags.get('weekly', []):
                    if tag in categories:
                        categories[tag].append(row)
                        matched = True
                        break
            elif isinstance(tags, list):
                for tag in tags:
                    if tag in categories:
                        categories[tag].append(row)
                        matched = True
                        break
            if not matched:
                uncategorized.append(row)
        if uncategorized:
            categories['宏观态势'].extend(uncategorized)

        total_sel = 0
        total_all = 0
        for cat in CATS:
            arts = categories[cat]
            if not arts:
                continue
            total_all += len(arts)
            sel = sum(1 for _, r in enumerate(arts) if st.session_state.wr_selections.get(str(r['id']), True))
            total_sel += sel

            with st.expander(f"{cat}  ({sel}/{len(arts)} 已选)", expanded=True):
                c1, c2 = st.columns([1, 10])
                with c1:
                    if st.button("全选", key=f"wr_sa_{cat}"):
                        for _, r in enumerate(arts):
                            st.session_state.wr_selections[str(r['id'])] = True
                        st.rerun()
                    if st.button("清空", key=f"wr_clr_{cat}"):
                        for _, r in enumerate(arts):
                            st.session_state.wr_selections[str(r['id'])] = False
                        st.rerun()

                for _, row in enumerate(arts):
                    aid = str(row['id'])
                    ptype = row.get('page_type', '')
                    badge = f"[{ptype}] " if ptype else ""
                    date_str = str(row.get('liangke_date', ''))[:10]
                    label = f"{badge}{date_str} | {str(row['title'])[:120]}"
                    checked = st.checkbox(label, value=st.session_state.wr_selections.get(aid, True), key=f"wr_cb_{aid}")
                    st.session_state.wr_selections[aid] = checked

        st.markdown(f"**总计已选：{total_sel} / {total_all}**")

    st.markdown("---")
    st.markdown("### 招投标数据")
    tender_file = st.file_uploader("上传招投标 Excel", type=['xlsx', 'xls'], key="wr_tender")

    st.markdown("---")
    PATENT_CATS = ['低温环境系统', '超导量子测控技术', '量子软件与算法', '量子算力网', '量子科技长三角产业创新中心']
    st.markdown("### 专利数据（一次上传多个 Excel，按文件名自动归入板块）")
    patent_raw = st.file_uploader("上传专利 Excel", type=['xlsx', 'xls'], accept_multiple_files=True, key="wr_pat_all",
                                  help="文件名需包含板块名，如：低温环境系统_0516.xlsx")
    patent_files = {}
    if patent_raw:
        for pf in patent_raw:
            name = pf.name
            matched = False
            for cat in PATENT_CATS:
                if cat in name:
                    patent_files[cat] = pf
                    st.caption(f"  + {name} -> {cat}")
                    matched = True
                    break
            if not matched:
                st.warning(f"  ! {name} not matched")

    if 'wr_result' not in st.session_state:
        st.session_state.wr_result = None
    if 'wr_pdf_path' not in st.session_state:
        st.session_state.wr_pdf_path = None

    btn_col1, btn_col2 = st.columns([2, 5])
    with btn_col1:
        gen_clicked = st.button("Generate PDF", type="primary", key="wr_gen")
    with btn_col2:
        if st.session_state.wr_result == 'success' and st.session_state.wr_pdf_path:
            with open(st.session_state.wr_pdf_path, 'rb') as f:
                st.download_button(
                    label="Download PDF", data=f,
                    file_name=os.path.basename(st.session_state.wr_pdf_path),
                    mime="application/pdf",
                )
            st.success("Success!")
        elif st.session_state.wr_result == 'error':
            st.error("PDF not found.")

    if gen_clicked:
        if not st.session_state.wr_loaded:
            st.error("Click 'Load Preview' first")
        else:
            with st.spinner("Generating..."):
                import subprocess
                temp_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'weekly_temp')
                os.makedirs(temp_dir, exist_ok=True)

                selected_ids = [aid for aid, v in st.session_state.wr_selections.items() if v]
                if not selected_ids:
                    st.error("No articles selected!")
                    return

                cmd = [
                    sys.executable,
                    os.path.join(os.path.dirname(os.path.abspath(__file__)), 'generate_weekly_report.py'),
                    '--start', start_date.strftime('%Y-%m-%d'),
                    '--end', end_date.strftime('%Y-%m-%d'),
                    '--issue', issue_no,
                    '--conf-month', str(conf_month),
                    '--page-types', ','.join(wr_page_types),
                    '--selected-ids', ','.join(selected_ids),
                ]

                if tender_file:
                    tender_path = os.path.join(temp_dir, 'tenders.xlsx')
                    with open(tender_path, 'wb') as f:
                        f.write(tender_file.getvalue())
                    cmd.extend(['--tender-excel', tender_path])

                for cat in PATENT_CATS:
                    pf = patent_files.get(cat)
                    if pf:
                        df = pd.read_excel(pf)
                        patents = parse_patent_excel(df, file_bytes=pf.getvalue())
                        if patents:
                            patents = filter_patents_by_date(patents, start_date, end_date)
                            filtered = filter_patents_llm(patents, cat) if len(patents) > 8 else patents
                            st.caption(f"{cat}: {len(patents)} raw -> {len(filtered)} selected")
                            pd.DataFrame(filtered).to_excel(
                                os.path.join(temp_dir, f'patents_{cat}.xlsx'), index=False, engine='openpyxl')
                            cmd.extend(['--patent-excel', os.path.join(temp_dir, f'patents_{cat}.xlsx')])

                result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8', errors='replace')
                st.code(result.stdout)

                if result.returncode != 0:
                    st.error("Failed")
                    st.code(result.stderr)
                    st.session_state.wr_result = 'error'
                    st.session_state.wr_pdf_path = None
                else:
                    pdf_path = os.path.join(
                        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        'weekly_output',
                        f'量子行业每周新闻洞察_第{issue_no}期.pdf'
                    )
                    if os.path.exists(pdf_path):
                        st.session_state.wr_result = 'success'
                        st.session_state.wr_pdf_path = pdf_path
                        st.rerun()
                    else:
                        st.error("PDF not found")
                        st.session_state.wr_result = 'error'
                        st.session_state.wr_pdf_path = None


# ------------------------------------------------------------------
# Page: Conferences
# ------------------------------------------------------------------

# ------------------------------------------------------------------
# Page: Report Alerts
# ------------------------------------------------------------------

def page_report_alerts():
    st.header("📢 报告提醒")

    # Prominent banner for recent reports
    import json, os
    from datetime import datetime, timedelta
    alert_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'report_alerts.json')
    if os.path.exists(alert_path):
        with open(alert_path, 'r', encoding='utf-8') as f:
            all_alerts = json.load(f)
        cutoff = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')
        recent = [a for a in all_alerts if a.get('date', '') >= cutoff]
        if recent:
            st.success(f"🆕 近 7 天发现 **{len(recent)}** 份新报告/白皮书/路线图")

    tab1, tab2 = st.tabs(["📰 新闻中发现的报告", "📚 光子盒报告"])

    with tab1:
        alert_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'report_alerts.json')
        if not os.path.exists(alert_path):
            st.info("暂无报告提醒。每日抓取后运行 scan_reports.py 即可自动扫描。")
        else:
            with open(alert_path, 'r', encoding='utf-8') as f:
                alerts = json.load(f)
            if not alerts:
                st.info("暂无报告提醒。")
            else:
                st.caption(f"共 {len(alerts)} 条（LLM 从每日新闻中自动识别）")
                for a in alerts:
                    with st.container():
                        st.markdown(f"### {a.get('report_name', '未知')}")
                        st.caption(f"📅 {a.get('date', '')}  |  🏛️ {a.get('publisher', '')}")
                        st.markdown(a.get('note', ''))
                        col1, col2 = st.columns(2)
                        src_url = a.get('source_article_url', '')
                        if src_url:
                            col1.link_button("📰 来源文章", src_url)
                        if a.get('url'):
                            col2.link_button("📄 报告下载", a['url'])
                        st.caption(f"来源新闻：{a.get('title', '')[:80]}")
                        st.markdown("---")

    with tab2:
        resources_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'resources.db')
        if not os.path.exists(resources_path):
            st.info("暂无光子盒报告。运行 scrape_reports_photon.py 抓取。")
        else:
            import sqlite3
            conn = sqlite3.connect(resources_path)
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute("SELECT * FROM reports ORDER BY publish_date DESC")
            rows = [dict(r) for r in c.fetchall()]
            conn.close()
            if not rows:
                st.info("暂无光子盒报告。")
            else:
                st.caption(f"共 {len(rows)} 份年度行业报告")
                for r in rows:
                    with st.container():
                        st.markdown(f"### {r.get('title', '')}")
                        st.caption(f"🏛️ {r.get('publisher', '')}  |  📅 {r.get('publish_date', '')}")
                        dl = r.get('download_url', '')
                        if dl:
                            st.link_button("📥 下载 PDF", dl)
                        st.markdown("---")

    with tab2:
        resources_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'resources.db')
        if not os.path.exists(resources_path):
            st.info("暂无外部资源。运行 scrape_reports_photon.py 可抓取光子盒等行业报告。")
        else:
            import sqlite3
            conn = sqlite3.connect(resources_path)
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute("SELECT * FROM reports ORDER BY publish_date DESC")
            rows = [dict(r) for r in c.fetchall()]
            conn.close()
            if not rows:
                st.info("暂无外部资源。")
            else:
                st.caption(f"共 {len(rows)} 份行业报告（外部网站抓取）")
                for r in rows:
                    with st.container():
                        st.markdown(f"### {r.get('title', '')}")
                        st.caption(f"🏛️ {r.get('publisher', '')}  |  📅 {r.get('publish_date', '')}")
                        link = r.get('download_url') or r.get('source_url', '')
                        if link:
                            st.link_button("🔗 查看", link)
                        if r.get('abstract'):
                            st.markdown(r['abstract'])
                        st.markdown("---")


def page_knowledge_graph():
    """知识图谱页面：pyvis 交互式实体关系网络."""
    import json, os
    from collections import defaultdict
    from pyvis.network import Network

    GRAPH_PATH = 'D:/Claude_code/knowledge_graph/knowledge_graph.json'

    st.title('量子科技情报知识图谱')
    st.caption('从 11,000+ 篇文章中抽取的实体和关系网络')

    if not os.path.exists(GRAPH_PATH):
        st.warning('图谱文件不存在，请先运行 knowledge_graph/build_graph.py')
        return

    with open(GRAPH_PATH, 'r', encoding='utf-8') as f:
        data = json.load(f)

    stats = data['meta']['stats']

    # Sidebar filters
    st.sidebar.markdown('---')
    st.sidebar.subheader('图谱筛选')

    node_types = defaultdict(int)
    for n in data['nodes']:
        node_types[n['type']] += 1

    entity_type_options = [t for t in sorted(node_types.keys()) if t != 'article']
    selected_types = st.sidebar.multiselect(
        '实体类型',
        options=entity_type_options,
        default=[t for t in ['institution', 'technology', 'product', 'topic', 'person'] if t in entity_type_options],
    )
    min_count = st.sidebar.slider('最少文章数', 1, 200, 1)
    search = st.sidebar.text_input('搜索实体', placeholder='输入名称...')

    st.sidebar.metric('总节点', stats['total_nodes'])
    st.sidebar.metric('总边', stats['total_edges'])
    st.sidebar.caption(f'文章: {stats["node_types"].get("article", 0)} 篇')
    st.sidebar.caption(f'每 6 小时自动更新')

    # Filter nodes
    COLORS = {
        'institution': '#4e79a7', 'technology': '#f28e2b', 'product': '#e15759',
        'topic': '#76b7b2', 'person': '#59a14f',
    }

    filtered = [n for n in data['nodes'] if n['type'] in selected_types and n['count'] >= min_count
                and (not search or search.lower() in n['id'].lower())]
    entity_ids = {n['id'] for n in filtered}

    # When searching and results are sparse, auto-expand to include neighbors
    if search and 0 < len(filtered) <= 10:
        neighbor_ids = set()
        for e in data['edges']:
            if e['source'] in entity_ids and e['target'] not in entity_ids:
                # Only add neighbor if its type is in selected_types
                tgt_node = next((n for n in data['nodes'] if n['id'] == e['target']), None)
                if tgt_node and tgt_node['type'] in selected_types:
                    neighbor_ids.add(e['target'])
            elif e['target'] in entity_ids and e['source'] not in entity_ids:
                src_node = next((n for n in data['nodes'] if n['id'] == e['source']), None)
                if src_node and src_node['type'] in selected_types:
                    neighbor_ids.add(e['source'])
        # Add neighbor nodes
        for n in data['nodes']:
            if n['id'] in neighbor_ids:
                filtered.append(n)
        entity_ids.update(neighbor_ids)

    edges_filtered = [e for e in data['edges'] if e['source'] in entity_ids and e['target'] in entity_ids]

    # Relation type translations
    REL_CN = {
        'PARTNERS_WITH': '合作', 'ACQUIRES': '收购', 'SUPPLIES_TO': '供应',
        'COMPETES_WITH': '竞争', 'MENTIONS': '提及', 'PUBLISHED_BY': '发布',
        'COVERS_TOPIC': '覆盖主题', 'USES_TECH': '使用技术', 'RELEASES': '发布产品',
        'WORKS_AT': '任职',
    }

    st.subheader(f'实体关系网络 ({len(filtered)} 节点, {len(edges_filtered)} 边)')
    st.caption('点击边查看关联新闻')

    net = Network(height='650px', width='100%', directed=True, notebook=False)
    net.set_options("""
    {
      "physics": {"barnesHut": {"gravitationalConstant": -3000, "centralGravity": 0.3,
                  "springLength": 200, "springConstant": 0.04}, "minVelocity": 0.75},
      "interaction": {"hover": true, "tooltipDelay": 100},
      "edges": {"smooth": false}
    }
    """)

    # Build edge data JS - embed article titles for click display
    edge_data_js = {}
    for e in edges_filtered:
        eid = f"{e['source']}|||{e['target']}|||{e['relation']}"
        arts = []
        try:
            raw = e.get('articles', '')
            arts = json.loads(raw) if isinstance(raw, str) else (raw or [])
        except (json.JSONDecodeError, TypeError):
            pass
        edge_data_js[eid] = {
            'source': e['source'], 'target': e['target'],
            'relation': REL_CN.get(e['relation'], e['relation']),
            'count': e.get('count', 0),
            'articles': arts[:20],
            'reason': e.get('reason', ''),
        }

    max_count = max((n['count'] for n in filtered), default=1)
    for n in filtered:
        size = 10 + 30 * (n['count'] / max_count)
        label = n['id'] if n['type'] not in ('institution', 'product') else f'{n["id"]}\n({n["count"]}篇)'
        net.add_node(n['id'], label=label, title=f'{n["id"]}\n类型: {n["type"]}\n文章数: {n["count"]}',
                     color=COLORS.get(n['type'], '#999'), size=size)

    DIRECTIONAL = {'收购', '供应', '任职', '发布产品'}
    for e in edges_filtered:
        cn_rel = REL_CN.get(e['relation'], e['relation'])
        year_info = ''
        years = e.get('years', [])
        if years and len(years) > 0:
            year_info = f' ({years[0]}' + (f'~{years[-1]}' if len(years) > 1 else '') + ')'
        tooltip = f"{cn_rel}: {e['source']} ↔ {e['target']}{year_info} ({e.get('count', '?')}篇)"
        arrows = 'to' if cn_rel in DIRECTIONAL else ''
        net.add_edge(e['source'], e['target'], title=tooltip, label=f"{cn_rel}{year_info}", arrows=arrows, physics=True)

    html_path = 'D:/Claude_code/knowledge_graph/graph_temp.html'
    net.save_graph(html_path)

    # Inject JS for click-to-show-articles + right panel CSS
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()

    edge_json = json.dumps(edge_data_js, ensure_ascii=False)
    inject_js = f"""
    <style>
    #edge-panel {{ position:fixed; right:10px; top:60px; width:320px; max-height:80vh;
      background:#fff; border:1px solid #ddd; border-radius:8px; padding:15px;
      box-shadow:0 4px 12px rgba(0,0,0,.15); overflow-y:auto; z-index:9999; display:none;
      font-family:-apple-system,BlinkMacSystemFont,sans-serif; font-size:13px; }}
    #edge-panel h4 {{ margin:0 0 8px; font-size:15px; color:#333; }}
    #edge-panel .rel-tag {{ display:inline-block; background:#4e79a7; color:#fff;
      padding:2px 8px; border-radius:4px; font-size:12px; margin-bottom:10px; }}
    #edge-panel .art-item {{ padding:6px 0; border-bottom:1px solid #eee; color:#555; line-height:1.4; }}
    #edge-panel .art-item:last-child {{ border-bottom:none; }}
    #edge-panel .close-btn {{ float:right; cursor:pointer; font-size:18px; color:#999; }}
    #edge-panel .close-btn:hover {{ color:#333; }}
    </style>
    <div id="edge-panel"><span class="close-btn" onclick="closePanel()">×</span><div id="edge-content"></div></div>
    <script>
    var edgeData = {edge_json};
    var network = null;
    function showPanel(edgeId) {{
        var data = edgeData[edgeId];
        if (!data) return;
        var panel = document.getElementById('edge-panel');
        var content = document.getElementById('edge-content');
        var html = '<h4>'+data.relation+': '+data.source+' ↔ '+data.target+'</h4>';
        html += '<span class="rel-tag">'+data.count+'篇</span>';
        if (data.reason) html += '<p style="color:#888;font-size:12px;margin-top:6px;">'+data.reason+'</p>';
        if (data.articles && data.articles.length > 0) {{
            html += '<div style="margin-top:10px;">';
            for (var i=0; i<data.articles.length; i++) {{
                html += '<div class="art-item">• '+data.articles[i]+'</div>';
            }}
            html += '</div>';
        }} else {{
            html += '<p style="color:#aaa;margin-top:10px;">暂无关联文章</p>';
        }}
        content.innerHTML = html;
        panel.style.display = 'block';
    }}
    function closePanel() {{ document.getElementById('edge-panel').style.display = 'none'; }}
    </script>
    """

    # Insert edge data + CSS before </body>, and hook selectEdge event
    html = html.replace('</body>', inject_js + '</body>')
    # Add edge click handler right after network creation
    edge_handler = """
    network.on("selectEdge", function(params) {
        if (params.edges.length > 0) {
            var edge = network.body.data.edges.get(params.edges[0]);
            if (edge) {
                var prefix = edge.from + "|||" + edge.to + "|||";
                for (var k in edgeData) {
                    if (k.indexOf(prefix) === 0) { showPanel(k); break; }
                }
            }
        }
    });
    network.on("deselectEdge", function() { closePanel(); });
    """
    html = html.replace('network = new vis.Network(container, data, options);',
                        'network = new vis.Network(container, data, options);\n' + edge_handler)
    st.components.v1.html(html, height=750, scrolling=True)

    # Entity table
    st.markdown('---')
    st.subheader('实体列表')
    rows = []
    for n in sorted(filtered, key=lambda x: -x['count']):
        connected = set()
        for e in edges_filtered:
            if e['source'] == n['id']:
                connected.add(f'{e["target"]}({e["relation"]})')
            elif e['target'] == n['id']:
                connected.add(f'{e["source"]}({e["relation"]})')
        rows.append({'名称': n['id'], '类型': n['type'], '文章数': n['count'], '关联': ', '.join(sorted(connected)[:5])})
    st.dataframe(rows, use_container_width=True, hide_index=True)


def count_new_reports():
    """Count reports added in the last 7 days."""
    import json, os
    from datetime import datetime, timedelta
    alert_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'report_alerts.json')
    if not os.path.exists(alert_path):
        return 0
    with open(alert_path, 'r', encoding='utf-8') as f:
        alerts = json.load(f)
    cutoff = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')
    return sum(1 for a in alerts if a.get('date', '') >= cutoff)


def main():
    st.set_page_config(page_title="量子科技情报", page_icon="📰", layout="wide")
    st.title("📰 量子科技情报")

    # Report alert badge
    new_reports = count_new_reports()
    nav_items = ["每日资讯", "周报生成", "会议信息"]
    if new_reports > 0:
        nav_items.append(f"报告提醒 🔴{new_reports}")
    else:
        nav_items.append("报告提醒")
    nav_items.append("知识图谱")

    page_raw = st.sidebar.radio("导航", nav_items)
    # Strip badge for page routing
    page = page_raw.split(" 🔴")[0] if " 🔴" in page_raw else page_raw

    if page == "每日资讯":
        page_daily_news()
    elif page == "周报生成":
        page_weekly_report()
    elif page == "会议信息":
        page_conferences()
    elif page == "报告提醒":
        page_report_alerts()
    elif page == "知识图谱":
        page_knowledge_graph()


if __name__ == '__main__':
    main()
